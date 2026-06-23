"""
Extractor universal de documentos para el clasificador NCM.

Acepta cualquier formato común que el cliente pueda mandar (PDF, Word, Excel,
HTML, texto plano, imágenes incluyendo HEIC de iPhone, BMP, TIFF) y devuelve:

    - texto consolidado (string) — para los documentos
    - imágenes (list[(bytes, mime_type)]) — listas para enviar a Claude vision

Los PDFs escaneados (sin texto extraíble) se convierten automáticamente a
imágenes de cada página y se incluyen en la lista de imágenes.
"""

from __future__ import annotations

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Claude vision soporta nativo: JPEG, PNG, GIF, WEBP.
# Cualquier otro lo convertimos a JPEG.
_MIME_CLAUDE_NATIVO = {"image/jpeg", "image/png", "image/gif", "image/webp"}

# Extensiones por categoría
EXTS_PDF = {"pdf"}
EXTS_DOCX = {"docx"}
EXTS_XLSX = {"xlsx", "xls"}
EXTS_TEXTO = {"txt", "md", "csv", "log"}
EXTS_HTML = {"html", "htm"}
EXTS_IMAGEN_NATIVA = {"jpg", "jpeg", "png", "webp", "gif"}
EXTS_IMAGEN_CONVERTIR = {"heic", "heif", "bmp", "tiff", "tif"}

# Para mostrar al usuario en el uploader de Streamlit
TODAS_EXTS = sorted(
    EXTS_PDF | EXTS_DOCX | EXTS_XLSX | EXTS_TEXTO | EXTS_HTML
    | EXTS_IMAGEN_NATIVA | EXTS_IMAGEN_CONVERTIR
)


def _ext(nombre: str) -> str:
    return nombre.lower().rsplit(".", 1)[-1] if "." in nombre else ""


# ── PDF ─────────────────────────────────────────────────────────────

def _extraer_texto_pdf(pdf_bytes: bytes) -> str:
    try:
        import pdfplumber
        partes = []
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    partes.append(t)
        return "\n".join(partes)
    except Exception as e:
        logger.warning("Error extrayendo texto de PDF: %s", e)
        return ""


def _pdf_a_imagenes(pdf_bytes: bytes, max_paginas: int = 10) -> list[tuple[bytes, str]]:
    """Renderiza cada página del PDF como JPEG. Útil para PDFs escaneados.

    Limita a max_paginas para no mandarle 200 páginas a Claude.
    """
    try:
        import fitz  # PyMuPDF
        imagenes: list[tuple[bytes, str]] = []
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            if i >= max_paginas:
                break
            # 144 DPI da buena calidad para Claude vision sin saturar tokens
            pix = page.get_pixmap(dpi=144)
            jpeg_bytes = pix.tobytes("jpeg")
            imagenes.append((jpeg_bytes, "image/jpeg"))
        doc.close()
        return imagenes
    except Exception as e:
        logger.warning("Error convirtiendo PDF a imágenes: %s", e)
        return []


# ── Word / Excel / HTML / texto ─────────────────────────────────────

def _extraer_texto_docx(docx_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        partes = [p.text for p in doc.paragraphs if p.text.strip()]
        # Incluir tablas
        for table in doc.tables:
            for row in table.rows:
                fila = " | ".join(cell.text.strip() for cell in row.cells)
                if fila.strip(" |"):
                    partes.append(fila)
        return "\n".join(partes)
    except Exception as e:
        logger.warning("Error extrayendo texto de DOCX: %s", e)
        return ""


def _extraer_texto_xlsx(xlsx_bytes: bytes) -> str:
    try:
        import pandas as pd
        sheets = pd.read_excel(io.BytesIO(xlsx_bytes), sheet_name=None)
        partes = []
        for sheet_name, df in sheets.items():
            partes.append(f"### Hoja: {sheet_name}")
            partes.append(df.to_csv(index=False, sep="\t"))
        return "\n".join(partes)
    except Exception as e:
        logger.warning("Error extrayendo texto de XLSX: %s", e)
        return ""


def _extraer_texto_html(html_bytes: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_bytes, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except Exception as e:
        logger.warning("Error extrayendo texto de HTML: %s", e)
        try:
            return html_bytes.decode("utf-8", errors="replace")
        except Exception:
            return ""


def _extraer_texto_plano(bytes_: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return bytes_.decode(enc)
        except UnicodeDecodeError:
            continue
    return bytes_.decode("utf-8", errors="replace")


# ── Imágenes ────────────────────────────────────────────────────────

_HEIF_REGISTRADO = False


def _registrar_heif():
    """pillow-heif se registra una sola vez en el runtime."""
    global _HEIF_REGISTRADO
    if _HEIF_REGISTRADO:
        return
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
        _HEIF_REGISTRADO = True
    except Exception as e:
        logger.warning("No se pudo registrar HEIF: %s", e)


def _convertir_imagen_a_jpeg(img_bytes: bytes, ext: str) -> Optional[bytes]:
    """Convierte HEIC/BMP/TIFF a JPEG para que Claude la acepte."""
    try:
        from PIL import Image
        if ext in ("heic", "heif"):
            _registrar_heif()
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode in ("RGBA", "LA", "P"):
            # Aplanar transparencia sobre blanco
            fondo = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            fondo.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
            img = fondo
        elif img.mode != "RGB":
            img = img.convert("RGB")
        # Limitar tamaño máximo para no exceder límites de Claude
        max_lado = 2048
        if max(img.size) > max_lado:
            img.thumbnail((max_lado, max_lado))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        return buf.getvalue()
    except Exception as e:
        logger.warning("No se pudo convertir imagen %s: %s", ext, e)
        return None


# ── API pública ─────────────────────────────────────────────────────

class ResultadoExtraccion:
    """Resultado de procesar uno o más archivos."""

    def __init__(self):
        self.texto: str = ""
        self.imagenes: list[tuple[bytes, str]] = []
        self.warnings: list[str] = []
        self.archivos_procesados: int = 0

    def __repr__(self):
        return (
            f"<Extraccion: {self.archivos_procesados} archivos, "
            f"{len(self.texto)} chars texto, "
            f"{len(self.imagenes)} imágenes, "
            f"{len(self.warnings)} warnings>"
        )


def procesar_archivos(archivos) -> ResultadoExtraccion:
    """
    Procesa una lista de UploadedFile de Streamlit (o tuplas (nombre, bytes, mime))
    y devuelve un ResultadoExtraccion con todo lo extraído.

    Acepta:
        - Documentos: PDF, DOCX, XLSX, XLS, TXT, MD, CSV, LOG, HTML, HTM
        - Imágenes nativas Claude: JPG, JPEG, PNG, WEBP, GIF
        - Imágenes a convertir: HEIC, HEIF, BMP, TIFF, TIF
    """
    resultado = ResultadoExtraccion()
    if not archivos:
        return resultado

    textos: list[str] = []

    for archivo in archivos:
        # Normalizar entrada (puede ser UploadedFile o tupla)
        if hasattr(archivo, "name") and hasattr(archivo, "read"):
            nombre = archivo.name
            bytes_ = archivo.read()
            mime = getattr(archivo, "type", None)
        else:
            nombre, bytes_, mime = archivo

        ext = _ext(nombre)
        resultado.archivos_procesados += 1

        try:
            if ext in EXTS_PDF:
                texto = _extraer_texto_pdf(bytes_)
                if texto and texto.strip():
                    textos.append(f"--- {nombre} ---\n{texto}")
                else:
                    # PDF escaneado o sin texto → renderizar como imágenes
                    imgs = _pdf_a_imagenes(bytes_)
                    if imgs:
                        resultado.imagenes.extend(imgs)
                        resultado.warnings.append(
                            f"{nombre}: PDF sin texto extraíble → "
                            f"se mandaron {len(imgs)} página(s) como imagen a Claude."
                        )
                    else:
                        resultado.warnings.append(f"{nombre}: no se pudo procesar el PDF.")

            elif ext in EXTS_DOCX:
                texto = _extraer_texto_docx(bytes_)
                if texto.strip():
                    textos.append(f"--- {nombre} ---\n{texto}")
                else:
                    resultado.warnings.append(f"{nombre}: Word vacío o ilegible.")

            elif ext in EXTS_XLSX:
                texto = _extraer_texto_xlsx(bytes_)
                if texto.strip():
                    textos.append(f"--- {nombre} ---\n{texto}")
                else:
                    resultado.warnings.append(f"{nombre}: Excel vacío o ilegible.")

            elif ext in EXTS_HTML:
                texto = _extraer_texto_html(bytes_)
                if texto.strip():
                    textos.append(f"--- {nombre} ---\n{texto}")

            elif ext in EXTS_TEXTO:
                texto = _extraer_texto_plano(bytes_)
                if texto.strip():
                    textos.append(f"--- {nombre} ---\n{texto}")

            elif ext in EXTS_IMAGEN_NATIVA:
                m = mime or f"image/{'jpeg' if ext == 'jpg' else ext}"
                if m not in _MIME_CLAUDE_NATIVO:
                    m = f"image/{ext if ext != 'jpg' else 'jpeg'}"
                resultado.imagenes.append((bytes_, m))

            elif ext in EXTS_IMAGEN_CONVERTIR:
                jpeg = _convertir_imagen_a_jpeg(bytes_, ext)
                if jpeg:
                    resultado.imagenes.append((jpeg, "image/jpeg"))
                else:
                    resultado.warnings.append(
                        f"{nombre}: no se pudo convertir {ext.upper()} a JPEG."
                    )

            else:
                resultado.warnings.append(
                    f"{nombre}: formato .{ext} no soportado, ignorado."
                )

        except Exception as e:
            logger.exception("Error procesando %s", nombre)
            resultado.warnings.append(f"{nombre}: error al procesar — {e}")

    resultado.texto = "\n\n".join(textos)
    return resultado
